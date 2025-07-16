from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime, timedelta

def format_money(amount):
    """Formate un montant en format comptable"""
    return "{:,.2f}".format(amount).replace(",", " ").replace(".00", "")

def get_default_values(is_sarl_au=False):
    """Retourne les valeurs par défaut pour le template"""
    current_date = datetime.now()
    next_month = current_date + timedelta(days=30)
    previous_year = current_date.year - 1

    # Valeurs par défaut pour l'entreprise
    company_defaults = {
        'nom_entreprise': '[Nom de la Société]',
        'capital': format_money(100000),
        'adresse': '[Adresse complète du siège social]',
        'rc_number': '[Numéro RC]',
        'if_number': '[Numéro IF]',
        'ice_number': '[Numéro ICE]',
        'cnss_number': '[Numéro CNSS]',
    }

    # Valeurs par défaut pour l'AGO
    ago_defaults = {
        'date_ago': next_month.strftime("%d/%m/%Y"),
        'heure_ago': '10:00',
        'lieu_ago': 'siège social de la société',
        'exercice': str(previous_year),
    }

    # Résolutions par défaut
    resolutions_defaults = {
        'resolution1': f"""L'Assemblée Générale, après avoir entendu la lecture du rapport de gestion du Gérant et après examen des comptes de l'exercice clos le 31 décembre {previous_year}, approuve les comptes annuels arrêtés à cette date, tels qu'ils ont été présentés, ainsi que les opérations traduites dans ces comptes ou résumées dans ce rapport.

En conséquence, elle donne quitus entier et sans réserve au Gérant pour l'exécution de son mandat au cours dudit exercice.""",

        'resolution2': f"""L'Assemblée Générale décide d'affecter le résultat de l'exercice clos le 31 décembre {previous_year}, soit un bénéfice net comptable de [montant en chiffres] DH ([montant en lettres] dirhams), de la manière suivante :

- Réserve légale (5%) : [montant] DH
- Distribution de dividendes : [montant] DH
- Report à nouveau : [montant] DH""",

        'resolution3': """L'Assemblée Générale donne quitus entier, définitif et sans réserve au Gérant pour sa gestion au cours de l'exercice écoulé."""
    }

    # Associés par défaut
    if is_sarl_au:
        associes_defaults = {
            'associe_unique': {
                'nom': 'M. [Nom et Prénom]',
                'parts': '1000',
                'pourcentage': '100%'
            }
        }
    else:
        associes_defaults = {
            'associe1': {
                'nom': 'M. [Nom et Prénom Associé 1]',
                'parts': '600',
                'pourcentage': '60%'
            },
            'associe2': {
                'nom': 'M. [Nom et Prénom Associé 2]',
                'parts': '400',
                'pourcentage': '40%'
            }
        }

    return {
        'company': company_defaults,
        'ago': ago_defaults,
        'resolutions': resolutions_defaults,
        'associes': associes_defaults
    }

def create_template(output_path, is_sarl_au=False):
    doc = Document()
    defaults = get_default_values(is_sarl_au)

    # Configuration de la mise en page
    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # Titre
    title = doc.add_paragraph("PROCÈS-VERBAL DE L'ASSEMBLÉE GÉNÉRALE ORDINAIRE ANNUELLE")
    title_format = title.paragraph_format
    title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True

    # En-tête
    doc.add_paragraph(defaults['company']['nom_entreprise']).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Société à Responsabilité Limitée" + (" à Associé Unique" if is_sarl_au else "")).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Au capital de {defaults['company']['capital']} dirhams").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Siège social : {defaults['company']['adresse']}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"RC n° {defaults['company']['rc_number']} - IF n° {defaults['company']['if_number']}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"ICE n° {defaults['company']['ice_number']} - CNSS n° {defaults['company']['cnss_number']}").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Espacement
    doc.add_paragraph()

    # Corps du document
    intro = doc.add_paragraph()
    intro.add_run(f"Le {defaults['ago']['date_ago']} à {defaults['ago']['heure_ago']}, ")
    intro.add_run("l'associé unique " if is_sarl_au else "les associés ")
    intro.add_run(f"de la société {defaults['company']['nom_entreprise']}, société à responsabilité limitée")
    intro.add_run(" à associé unique " if is_sarl_au else " ")
    intro.add_run(f"au capital de {defaults['company']['capital']} dirhams, ")
    intro.add_run("s'est réuni" if is_sarl_au else "se sont réunis")
    intro.add_run(f" en Assemblée Générale Ordinaire Annuelle au {defaults['ago']['lieu_ago']}.")

    # Ordre du jour
    doc.add_paragraph()
    doc.add_heading("ORDRE DU JOUR", level=1)
    ordre_jour = [
        f"1. Rapport de gestion du Gérant sur l'exercice clos le 31 décembre {defaults['ago']['exercice']}",
        "2. Approbation des comptes de cet exercice et quitus au Gérant",
        "3. Affectation du résultat de l'exercice",
        "4. Questions diverses"
    ]
    for point in ordre_jour:
        doc.add_paragraph(point, style='List Number')

    doc.add_paragraph()

    # Présidence et bureau
    doc.add_paragraph(f"L'assemblée est présidée par M. [Nom du Président], Gérant de la société {defaults['company']['nom_entreprise']}.")

    doc.add_paragraph()

    # Liste des associés
    doc.add_heading("COMPOSITION DE L'ASSEMBLÉE", level=1)
    if is_sarl_au:
        assoc = defaults['associes']['associe_unique']
        doc.add_paragraph(f"L'associé unique, {assoc['nom']}, détenant la totalité des {assoc['parts']} parts sociales composant le capital social ({assoc['pourcentage']}), est présent.")
    else:
        doc.add_paragraph("Sont présents ou représentés les associés suivants :")
        for key, assoc in defaults['associes'].items():
            doc.add_paragraph(f"- {assoc['nom']}, détenant {assoc['parts']} parts sociales ({assoc['pourcentage']})")
        doc.add_paragraph(f"Soit un total de 1000 parts sociales représentant 100% du capital social.")

    doc.add_paragraph()

    # Documents présentés
    docs = doc.add_paragraph("Le Président dépose sur le bureau et met à la disposition ")
    docs.add_run("de l'associé unique" if is_sarl_au else "des associés")
    docs.add_run(" les documents suivants :")

    doc.add_paragraph(f"- Les comptes annuels (bilan, compte de résultat et annexe) de l'exercice clos le 31 décembre {defaults['ago']['exercice']}")
    doc.add_paragraph("- Le rapport de gestion établi par la Gérance")
    doc.add_paragraph("- Le texte des résolutions proposées")
    if not is_sarl_au:
        doc.add_paragraph("- La feuille de présence certifiée exacte")
        doc.add_paragraph("- Les pouvoirs des associés représentés")

    doc.add_paragraph()

    # Résolutions
    doc.add_heading("DÉLIBÉRATIONS", level=1)

    # Première résolution
    doc.add_heading("PREMIÈRE RÉSOLUTION - APPROBATION DES COMPTES", level=2)
    doc.add_paragraph(defaults['resolutions']['resolution1'])
    doc.add_paragraph("Cette résolution est adoptée à l'unanimité.")
    doc.add_paragraph()

    # Deuxième résolution
    doc.add_heading("DEUXIÈME RÉSOLUTION - AFFECTATION DU RÉSULTAT", level=2)
    doc.add_paragraph(defaults['resolutions']['resolution2'])
    doc.add_paragraph("Cette résolution est adoptée à l'unanimité.")
    doc.add_paragraph()

    # Troisième résolution
    doc.add_heading("TROISIÈME RÉSOLUTION - QUITUS AU GÉRANT", level=2)
    doc.add_paragraph(defaults['resolutions']['resolution3'])
    doc.add_paragraph("Cette résolution est adoptée à l'unanimité.")
    doc.add_paragraph()

    # Clôture
    doc.add_paragraph(f"L'ordre du jour étant épuisé et personne ne demandant plus la parole, la séance est levée à [heure de fin].")

    doc.add_paragraph()
    doc.add_paragraph("De tout ce que dessus, il a été dressé le présent procès-verbal qui a été signé par le Président.")

    # Signatures
    doc.add_paragraph()
    doc.add_paragraph()

    sig_section = doc.add_paragraph()
    sig_section.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig_section.add_run("Le Président").bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig.add_run("[Signature]")

    # Sauvegarde du document
    doc.save(output_path)

def main():
    base_dir = os.path.dirname(os.path.abspath(__file__))
    templates_dir = os.path.join(base_dir, 'static', 'templates')
    os.makedirs(templates_dir, exist_ok=True)

    # Création du template SARL
    create_template(
        os.path.join(templates_dir, 'modele_pv_ago_sarl.docx'),
        is_sarl_au=False
    )

    # Création du template SARL AU
    create_template(
        os.path.join(templates_dir, 'modele_pv_ago_sarl_au.docx'),
        is_sarl_au=True
    )

if __name__ == '__main__':
    main()
