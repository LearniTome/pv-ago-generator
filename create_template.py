from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_template(output_path):
    doc = Document()

    # Style par défaut
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # Titre
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run('[OBJET_AGO]').bold = True

    doc.add_paragraph()

    # Informations de base
    info = doc.add_paragraph()
    info.add_run('Date: [DATE_AGO]\n')
    info.add_run('Lieu: [LIEU_AGO]\n')
    info.add_run('Heure: [HEURE_AGO]\n')

    doc.add_paragraph()

    # En-tête du PV
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.add_run('PROCÈS-VERBAL DE L\'ASSEMBLÉE GÉNÉRALE ORDINAIRE\n').bold = True

    # Informations de la société
    company = doc.add_paragraph()
    company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    company.add_run('{{nom_entreprise}}\n')
    company.add_run('\nCapital social : {{capital}} Dirhams')
    company.add_run('\nSiège social : {{adresse}}\n\n')

    # Corps du document
    body = doc.add_paragraph()
    body.add_run('Le [DATE_AGO] à [HEURE_AGO], les associés de la société {{nom_entreprise}} ')
    body.add_run('se sont réunis en Assemblée Générale Ordinaire à [LIEU_AGO] sur convocation ')
    body.add_run('régulière du Gérant.')

    # Sauvegarde du document
    doc.save(output_path)

if __name__ == '__main__':
    # Créer les deux modèles
    create_template('static/templates/modele_pv_ago_sarl.docx')
    create_template('static/templates/modele_pv_ago_sarl_au.docx')
