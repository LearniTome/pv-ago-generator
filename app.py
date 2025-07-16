from flask import Flask, render_template, request, redirect, url_for, session, send_file, jsonify
import os
from werkzeug.security import generate_password_hash, check_password_hash
from datetime import datetime
import shutil
from pathlib import Path
from docx import Document
import uuid
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
import pandas as pd
import json
from datetime import datetime, timedelta
import sqlite3

app = Flask(__name__)
app.secret_key = 'votre_cle_secrete_a_modifier'

# Configuration des chemins
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATES_DIR = os.path.join(BASE_DIR, 'static', 'templates')
STATIC_DIR = os.path.join(BASE_DIR, 'static')
TEMPLATES = {
    'SARL': os.path.join(TEMPLATES_DIR, 'modele_pv_ago_sarl.docx'),
    'SARL AU': os.path.join(TEMPLATES_DIR, 'modele_pv_ago_sarl_au.docx')
}

# Create templates directory if it doesn't exist
os.makedirs(TEMPLATES_DIR, exist_ok=True)

# Ensure template files exist
default_template = os.path.join(TEMPLATES_DIR, 'modele_pv_ago.docx')
if os.path.exists(default_template):
    # If only the default template exists, use it for both types
    if not os.path.exists(TEMPLATES['SARL']):
        shutil.copy2(default_template, TEMPLATES['SARL'])
    if not os.path.exists(TEMPLATES['SARL AU']):
        shutil.copy2(default_template, TEMPLATES['SARL AU'])

OUTPUT_DIR = os.path.join(STATIC_DIR, 'output')
TEMP_DIR = os.path.join(STATIC_DIR, 'temp')
DB_PATH = os.path.join(BASE_DIR, 'database.db')

# Créer les répertoires s'ils n'existent pas
os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(TEMP_DIR, exist_ok=True)

# Ajouter le dossier static/temp aux fichiers statiques servis par Flask
app.static_folder = STATIC_DIR

# Initialisation de la base de données
def init_db():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()

    # Table pour l'historique des PV
    c.execute('''CREATE TABLE IF NOT EXISTS pv_history
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  username TEXT NOT NULL,
                  date_creation TEXT NOT NULL,
                  type_societe TEXT NOT NULL,
                  nom_entreprise TEXT NOT NULL,
                  filename TEXT NOT NULL)''')

    # Table pour les templates favoris
    c.execute('''CREATE TABLE IF NOT EXISTS favorite_templates
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  username TEXT NOT NULL,
                  template_name TEXT NOT NULL,
                  template_data TEXT NOT NULL)''')

    # Table pour les statistiques
    c.execute('''CREATE TABLE IF NOT EXISTS statistics
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  date TEXT NOT NULL,
                  type_societe TEXT NOT NULL,
                  action TEXT NOT NULL)''')

    # Table pour les associés
    c.execute('''CREATE TABLE IF NOT EXISTS associes
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  pv_id INTEGER NOT NULL,
                  nom TEXT NOT NULL,
                  prenom TEXT NOT NULL,
                  adresse TEXT NOT NULL,
                  cni TEXT NOT NULL,
                  cni_validite TEXT NOT NULL,
                  cni_lieu TEXT,
                  email TEXT,
                  telephone TEXT,
                  nombre_parts INTEGER NOT NULL,
                  pourcentage REAL NOT NULL,
                  FOREIGN KEY (pv_id) REFERENCES pv_history(id))''')

    # Table pour les documents
    c.execute('''CREATE TABLE IF NOT EXISTS documents
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  type_societe TEXT NOT NULL,
                  objet_ago TEXT NOT NULL,
                  date_ago DATE NOT NULL,
                  lieu_ago TEXT NOT NULL,
                  heure_ago TIME NOT NULL,
                  docx_filename TEXT NOT NULL,
                  pdf_filename TEXT NOT NULL,
                  date_creation TIMESTAMP DEFAULT CURRENT_TIMESTAMP)''')

    conn.commit()
    conn.close()

# Initialisation de la base de données au démarrage
init_db()

# Configuration de base
USERS = {
    'admin': generate_password_hash('motdepasse')
}

@app.route('/', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        if username in USERS and check_password_hash(USERS[username], password):
            session['username'] = username
            return redirect(url_for('formulaire'))
        else:
            return render_template('login.html', error='Identifiants invalides')
    return render_template('login.html')

@app.route('/formulaire')
def formulaire():
    if 'username' not in session:
        return redirect(url_for('login'))

    # Récupérer les valeurs par défaut
    default_values = get_default_values()

    return render_template('formulaire.html', default_values=default_values)

@app.route('/logout')
def logout():
    session.pop('username', None)
    return redirect(url_for('login'))

def save_pv_history(username, type_societe, nom_entreprise, filename):
    """Enregistre l'historique de génération d'un PV."""
    try:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute('''INSERT INTO pv_history
                    (username, date_creation, type_societe, nom_entreprise, filename)
                    VALUES (?, ?, ?, ?, ?)''',
                 (username,
                  datetime.now().strftime('%Y-%m-%d'),
                  type_societe,
                  nom_entreprise,
                  filename))

        # Enregistrement des statistiques
        c.execute('''INSERT INTO statistics
                    (date, type_societe, action)
                    VALUES (?, ?, ?)''',
                 (datetime.now().strftime('%Y-%m-%d'),
                  type_societe,
                  'generate_pv'))

        conn.commit()
        conn.close()
    except Exception as e:
        print(f"Erreur lors de l'enregistrement de l'historique : {str(e)}")

@app.route('/generer', methods=['POST'])
def generer():
    if 'username' not in session:
        return redirect(url_for('login'))

    conn = None
    try:
        # Récupération des données du formulaire
        data = request.form.to_dict()

        # Détermination du type de société
        type_societe = data.get('type_societe', '')
        if not type_societe:
            type_societe = 'SARL'  # Type par défaut
            data['type_societe'] = type_societe

        if type_societe not in TEMPLATES:
            return jsonify({
                'success': False,
                'error': "Type de société non valide. Choisissez SARL ou SARL AU."
            }), 400

        # Récupérer les valeurs par défaut
        default_values = get_default_values(type_societe)

        # Compléter les données manquantes avec les valeurs par défaut
        for key, value in default_values.items():
            if key != 'associes' and (key not in data or not data[key]):
                data[key] = value

        try:
            # Formatage des données des associés
            associes = format_associes_data(data)
            if not associes:
                # Utiliser les associés par défaut si aucun n'est fourni
                associes = default_values['associes']
                data['associes'] = json.dumps(associes)
        except ValueError as e:
            return jsonify({
                'success': False,
                'error': str(e)
            }), 400

        # Vérification du total des parts
        total_pourcentage = sum(associe['pourcentage'] for associe in associes)
        if abs(total_pourcentage - 100) > 0.01:
            return jsonify({
                'success': False,
                'error': "Le total des parts doit être égal à 100%"
            }), 400

        # Vérification du nombre d'associés selon le type de société
        if type_societe == 'SARL AU' and len(associes) != 1:
            return jsonify({
                'success': False,
                'error': "Une SARL AU ne peut avoir qu'un seul associé."
            }), 400
        elif type_societe == 'SARL' and len(associes) < 2:
            return jsonify({
                'success': False,
                'error': "Une SARL doit avoir au moins deux associés."
            }), 400

        template_path = TEMPLATES[type_societe]
        if not os.path.exists(template_path):
            return jsonify({
                'success': False,
                'error': f"Le modèle de document pour {type_societe} n'existe pas."
            }), 400

        # Nettoyer les fichiers temporaires avant de générer
        cleanup_temp_files()

        # Générer le document
        result = generate_document(data, is_preview=False)

        if not result['success']:
            return jsonify({
                'success': False,
                'error': result.get('error', "Erreur lors de la génération du document")
            }), 500

        # Utiliser le nouveau format de nom de fichier
        output_filename = generate_filename(data, 'pdf')
        output_path = os.path.join(OUTPUT_DIR, output_filename)

        # Copier le fichier généré avec le nouveau nom
        shutil.copy2(result['pdf_path'], output_path)

        # Connexion à la base de données pour l'historique
        conn = sqlite3.connect(DB_PATH)
        save_pv_history(
            session['username'],
            type_societe,
            data.get('nom_entreprise', ''),
            output_filename
        )
        conn.commit()

        # Envoyer le fichier
        return send_file(
            output_path,
            as_attachment=True,
            download_name=output_filename,
            mimetype='application/pdf'
        )

    except Exception as e:
        app.logger.error(f"Erreur lors de la génération du document: {str(e)}")
        if conn:
            conn.rollback()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

    finally:
        if conn:
            conn.close()

@app.route('/convert_pdf', methods=['POST'])
def convert_pdf_route():
    if 'username' not in session:
        return redirect(url_for('login'))

    try:
        wordfile = request.form.get('wordfile')
        if not wordfile:
            raise ValueError("Nom de fichier Word non fourni")

        word_path = os.path.join(TEMP_DIR, wordfile)
        pdf_filename = wordfile.replace('.docx', '.pdf')
        pdf_path = os.path.join(TEMP_DIR, pdf_filename)

        # Utiliser la nouvelle fonction de conversion
        success = convert_to_pdf(word_path, pdf_path)
        if not success:
            raise Exception("Échec de la conversion PDF")

        # Envoi du fichier PDF
        return send_file(
            pdf_path,
            as_attachment=True,
            download_name=pdf_filename,
            mimetype='application/pdf'
        )

    except Exception as e:
        return jsonify({
            'success': False,
            'error': f"Erreur lors de la conversion en PDF : {str(e)}"
        }), 500

@app.route('/pv/<filename>')
def download_pv(filename):
    if 'username' not in session:
        return redirect(url_for('login'))
    return send_file(
        os.path.join(TEMP_DIR, filename),
        as_attachment=True
    )

# Fonction utilitaire pour nettoyer les fichiers temporaires
def cleanup_temp_files(keep_preview_id=None):
    """Nettoie tous les fichiers temporaires sauf la prévisualisation active si spécifiée."""
    import time
    from pathlib import Path

    def try_delete_file(file_path, max_attempts=3, delay=1):
        """Tente de supprimer un fichier plusieurs fois avec un délai entre les tentatives."""
        for attempt in range(max_attempts):
            try:
                if os.path.exists(file_path):
                    os.remove(file_path)
                return True
            except Exception as e:
                if attempt < max_attempts - 1:
                    time.sleep(delay)
                    continue
                else:
                    print(f"Impossible de supprimer {file_path} après {max_attempts} tentatives : {e}")
                    return False

    try:
        # Attendre un court instant pour s'assurer que Word a bien libéré les fichiers
        time.sleep(0.5)

        # Supprimer d'abord les fichiers temporaires de Word (.tmp)
        for file in Path(TEMP_DIR).glob('mso*.tmp'):
            try_delete_file(str(file))

        # Ensuite, nettoyer les autres fichiers temporaires
        for file in os.listdir(TEMP_DIR):
            file_path = os.path.join(TEMP_DIR, file)
            if os.path.isfile(file_path):
                # Si un ID de prévisualisation est fourni, ne pas supprimer ses fichiers
                if keep_preview_id and (f"preview_{keep_preview_id}.docx" in file or f"preview_{keep_preview_id}.pdf" in file):
                    continue
                try_delete_file(file_path)

    except Exception as e:
        print(f"Erreur lors du nettoyage des fichiers temporaires : {e}")

# Fonction unifiée pour générer un document (preview ou final)
def generate_document(data, is_preview=True):
    try:
        # Créer un ID unique pour le document
        doc_id = str(uuid.uuid4())

        # Nettoyer les fichiers de prévisualisation existants si c'est une prévisualisation
        if is_preview:
            cleanup_temp_files(keep_preview_id=doc_id)

        # Déterminer le type de société et le template approprié
        type_societe = data.get('type_societe', 'SARL')
        template_path = TEMPLATES.get(type_societe)

        if not template_path or not os.path.exists(template_path):
            raise ValueError(f"Template non trouvé pour le type de société : {type_societe}")

        # S'assurer que les répertoires existent
        if is_preview:
            os.makedirs(TEMP_DIR, exist_ok=True)
        else:
            os.makedirs(OUTPUT_DIR, exist_ok=True)

        # Créer une copie du template
        doc = Document(template_path)

        # Préparer toutes les variables de remplacement
        variables = {
            'nom_entreprise': data.get('nom_entreprise', ''),
            'type_societe': type_societe,
            'capital': data.get('capital', ''),
            'adresse': data.get('adresse', ''),
            'date_assemblee': data.get('date_assemblee', ''),
            'lieu_assemblee': data.get('lieu_assemblee', ''),
            'objet_ago': data.get('objet_ago', ''),
            'heure_ago': data.get('heure_ago', '')
        }

        # Remplacer les variables dans le document
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                text = run.text
                # Remplacer les variables avec les deux formats
                for key, value in variables.items():
                    text = text.replace(f"{{{{{key}}}}}", str(value))
                    text = text.replace(f"[[{key}]]", str(value))
                run.text = text

        # Ajouter la section des associés si présente
        if 'associes' in data:
            try:
                associes = json.loads(data['associes']) if isinstance(data['associes'], str) else data['associes']
                generate_associes_section(doc, associes)
            except json.JSONDecodeError:
                app.logger.warning("Impossible de décoder les données des associés")

        # Déterminer le chemin de sauvegarde
        if is_preview:
            filename = f"preview_{doc_id}.docx"
            output_dir = TEMP_DIR
        else:
            date_str = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"pv_{date_str}_{doc_id[:6]}.docx"
            output_dir = OUTPUT_DIR

        # Sauvegarder le document
        output_path = os.path.join(output_dir, filename)
        doc.save(output_path)

        # Convertir en PDF
        pdf_filename = os.path.splitext(filename)[0] + '.pdf'
        pdf_path = os.path.join(output_dir, pdf_filename)
        convert_to_pdf(output_path, pdf_path)

        # Enregistrer dans la base de données si ce n'est pas une prévisualisation
        if not is_preview:
            save_to_database(data, filename, pdf_filename)

        return {
            'success': True,
            'preview_id': doc_id,
            'filename': pdf_filename,
            'docx_path': output_path,
            'pdf_path': pdf_path
        }

    except Exception as e:
        app.logger.error(f"Erreur lors de la génération du document : {e}")
        return {'success': False, 'error': str(e)}

@app.route('/preview', methods=['POST'])
def preview_document():
    try:
        data = request.form.to_dict()
        current_step = data.get('current_step', '1')  # Par défaut étape 1

        # Si ce n'est pas la dernière étape (étape 3), on retourne juste un succès sans générer de fichier
        if current_step != '3':
            return jsonify({
                'success': True,
                'message': 'Validation réussie',
                'preview_skipped': True
            })

        # Uniquement générer le document dans la dernière étape
        result = generate_document(data, is_preview=True)

        if result['success']:
            return jsonify({
                'success': True,
                'preview_id': result['preview_id']
            })
        else:
            return jsonify({
                'success': False,
                'error': result['error']
            })
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        })

@app.route('/generate', methods=['POST'])
def generate_final_document():
    try:
        app.logger.info("Début de la génération du document final")
        data = request.form.to_dict()
        app.logger.info(f"Données reçues: {data}")

        # Validation des données requises
        required_fields = ['objet_ago', 'date_ago', 'lieu_ago', 'heure_ago']
        missing_fields = [field for field in required_fields if not data.get(field)]

        if missing_fields:
            raise ValueError(f"Champs obligatoires manquants: {', '.join(missing_fields)}")

        # Générer le document
        result = generate_document(data, is_preview=False)
        app.logger.info(f"Résultat de la génération: {result}")

        if result['success']:
            response_data = {
                'success': True,
                'filename': result['filename'],
                'download_url': url_for('get_document', filename=result['filename'], _external=True)
            }
            app.logger.info(f"Réponse préparée: {response_data}")
            return jsonify(response_data)
        else:
            raise ValueError(result.get('error', 'Erreur inconnue lors de la génération'))

    except Exception as e:
        app.logger.error(f"Erreur lors de la génération: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500  # Retourner un code 500 pour les erreurs serveur

@app.route('/get_preview/<preview_id>')
def get_preview(preview_id):
    try:
        # Chercher le fichier PDF de prévisualisation
        preview_file = None
        for file in os.listdir(TEMP_DIR):
            if file.startswith(f"preview_{preview_id}") and file.endswith('.pdf'):
                preview_file = file
                break

        if preview_file:
            return send_file(
                os.path.join(TEMP_DIR, preview_file),
                mimetype='application/pdf'
            )
        else:
            return "Prévisualisation non trouvée", 404
    except Exception as e:
        return str(e), 500

@app.route('/get_document/<filename>')
def get_document(filename):
    try:
        return send_file(
            os.path.join(OUTPUT_DIR, filename),
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        return str(e), 500

# Route pour sauvegarder un template favori
@app.route('/save_favorite', methods=['POST'])
def save_favorite():
    if 'username' not in session:
        return jsonify({'error': 'Non authentifié'}), 401
    try:
        data = request.json
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute('INSERT INTO favorite_templates (username, template_name, template_data) VALUES (?, ?, ?)',
                 (session['username'], data['name'], json.dumps(data['template'])))
        conn.commit()
        conn.close()
        return jsonify({'message': 'Template sauvegardé avec succès'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# Route pour récupérer les templates favoris
@app.route('/get_favorites')
def get_favorites():
    if 'username' not in session:
        return jsonify({'error': 'Non authentifié'}), 401
    try:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute('SELECT template_name, template_data FROM favorite_templates WHERE username = ?',
                 (session['username'],))
        favorites = [{'name': row[0], 'template': json.loads(row[1])} for row in c.fetchall()]
        conn.close()
        return jsonify({'favorites': favorites})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# Route for the associates
@app.route('/save_associes', methods=['POST'])
def save_associes_route():
    if 'username' not in session:
        return jsonify({'error': 'Non authentifié'}), 401
    try:
        data = request.json
        associes = data.get('associes', [])

        # Créer une entrée temporaire dans l'historique des PV pour les associés
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()

        try:
            # Insérer un PV temporaire
            c.execute('''INSERT INTO pv_history
                        (username, date_creation, type_societe, nom_entreprise, filename)
                        VALUES (?, ?, ?, ?, ?)''',
                     (session['username'],
                      datetime.now().strftime('%Y-%m-%d'),
                      'temp',
                      'temp',
                      'temp'))

            pv_id = c.lastrowid

            # Sauvegarder les associés
            for associe in associes:
                c.execute('''INSERT INTO associes
                            (pv_id, nom, prenom, adresse, cni, cni_validite, cni_lieu,
                             email, telephone, nombre_parts, pourcentage)
                            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)''',
                         (pv_id,
                          associe['nom'],
                          associe['prenom'],
                          associe['adresse'],
                          associe['cni'],
                          associe['cni_validite'],
                          associe.get('cni_lieu'),
                          associe.get('email'),
                          associe.get('telephone'),
                          associe['parts'],
                          associe['pourcentage']))

            conn.commit()
            return jsonify({
                'message': 'Associés sauvegardés avec succès',
                'pv_id': pv_id
            })

        except Exception as e:
            conn.rollback()
            raise e

        finally:
            conn.close()

    except Exception as e:
        return jsonify({'error': str(e)}), 500

# Route pour l'historique des PV
@app.route('/history')
def history():
    if 'username' not in session:
        return redirect(url_for('login'))
    try:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute('SELECT date_creation, type_societe, nom_entreprise, filename FROM pv_history WHERE username = ? ORDER BY date_creation DESC',
                 (session['username'],))
        history = c.fetchall()
        conn.close()
        return render_template('history.html', history=history)
    except Exception as e:
        return render_template('history.html', error=str(e))

# Route pour le tableau de bord des statistiques
@app.route('/dashboard')
def dashboard():
    if 'username' not in session:
        return redirect(url_for('login'))
    try:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()

        # Statistiques des 30 derniers jours
        thirty_days_ago = (datetime.now() - timedelta(days=30)).strftime('%Y-%m-%d')

        # PV par type de société
        c.execute('''SELECT type_societe, COUNT(*) as count
                    FROM pv_history
                    WHERE date_creation >= ?
                    GROUP BY type_societe''', (thirty_days_ago,))
        pv_stats = format_stats_data(c.fetchall())

        # Evolution journalière
        c.execute('''SELECT date_creation, COUNT(*) as count
                    FROM pv_history
                    WHERE date_creation >= ?
                    GROUP BY date_creation
                    ORDER BY date_creation''', (thirty_days_ago,))
        daily_stats = format_stats_data(c.fetchall())

        conn.close()
        return render_template('dashboard.html',
                             pv_stats=pv_stats,
                             daily_stats=daily_stats)
    except Exception as e:
        return render_template('dashboard.html', error=str(e))

# Route pour l'import Excel
@app.route('/import_excel', methods=['POST'])
def import_excel():
    if 'username' not in session:
        return jsonify({'error': 'Non authentifié'}), 401
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'Aucun fichier fourni'}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'Aucun fichier sélectionné'}), 400

        if not file.filename.endswith('.xlsx'):
            return jsonify({'error': 'Format de fichier non supporté'}), 400

        # Lecture du fichier Excel
        df = pd.read_excel(file)
        # Conversion en dictionnaire
        data = df.to_dict('records')[0] if not df.empty else {}

        return jsonify({'data': data})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

def format_stats_data(stats_data):
    """Formate les données statistiques pour les graphiques."""
    return [[str(item[0]), int(item[1])] for item in stats_data]



def format_associes_data(data):
    """Formate les données des associés pour le PV."""
    try:
        # Vérifier si les données des associés sont présentes
        if not any(key.startswith('associe_') for key in data):
            # Si aucune donnée d'associé n'est présente, retourner une liste vide
            return []

        associes = []
        noms = data.getlist('associe_nom[]') if hasattr(data, 'getlist') else data.get('associe_nom[]', [])
        prenoms = data.getlist('associe_prenom[]') if hasattr(data, 'getlist') else data.get('associe_prenom[]', [])
        adresses = data.getlist('associe_adresse[]') if hasattr(data, 'getlist') else data.get('associe_adresse[]', [])
        cnis = data.getlist('associe_cni[]') if hasattr(data, 'getlist') else data.get('associe_cni[]', [])
        cni_validites = data.getlist('associe_cni_validite[]') if hasattr(data, 'getlist') else data.get('associe_cni_validite[]', [])
        cni_lieux = data.getlist('associe_cni_lieu[]') if hasattr(data, 'getlist') else data.get('associe_cni_lieu[]', [])
        emails = data.getlist('associe_email[]') if hasattr(data, 'getlist') else data.get('associe_email[]', [])
        telephones = data.getlist('associe_telephone[]') if hasattr(data, 'getlist') else data.get('associe_telephone[]', [])
        parts = data.getlist('associe_parts[]') if hasattr(data, 'getlist') else data.get('associe_parts[]', [])
        pourcentages = data.getlist('associe_pourcentage[]') if hasattr(data, 'getlist') else data.get('associe_pourcentage[]', [])

        # Vérifier que toutes les listes ont la même longueur
        lists_length = len(noms)
        if not all(len(lst) == lists_length for lst in [prenoms, adresses, cnis, cni_validites, parts, pourcentages]):
            raise ValueError("Les données des associés sont incomplètes ou mal formatées")

        for i in range(lists_length):
            try:
                # S'assurer que les valeurs numériques sont valides
                part_value = int(parts[i]) if parts[i] else 0
                pourcentage_value = float(pourcentages[i].replace('%', '')) if pourcentages[i] else 0.0

                associe = {
                    'nom': noms[i].strip(),
                    'prenom': prenoms[i].strip(),
                    'adresse': adresses[i].strip(),
                    'cni': cnis[i].strip(),
                    'cni_validite': cni_validites[i].strip(),
                    'cni_lieu': cni_lieux[i].strip() if i < len(cni_lieux) and cni_lieux[i] else None,
                    'email': emails[i].strip() if i < len(emails) and emails[i] else None,
                    'telephone': telephones[i].strip() if i < len(telephones) and telephones[i] else None,
                    'parts': part_value,
                    'pourcentage': pourcentage_value
                }
                associes.append(associe)
            except (ValueError, IndexError) as e:
                app.logger.error(f"Erreur lors du traitement des données de l'associé {i+1}: {str(e)}")
                continue

        return associes

    except Exception as e:
        app.logger.error(f"Erreur lors du formatage des données des associés: {str(e)}")
        raise ValueError(f"Impossible de traiter les données des associés: {str(e)}")

def save_associes(pv_id, associes):
    """Enregistre les informations des associés dans la base de données."""
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()

    for associe in associes:
        c.execute('''INSERT INTO associes
                    (pv_id, nom, prenom, adresse, cni, cni_validite, cni_lieu,
                     email, telephone, nombre_parts, pourcentage)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)''',
                 (pv_id, associe['nom'], associe['prenom'], associe['adresse'],
                  associe['cni'], associe['cni_validite'], associe['cni_lieu'],
                  associe['email'], associe['telephone'], associe['parts'],
                  associe['pourcentage']))

    conn.commit()
    conn.close()

def generate_associes_section(doc, associes):
    """Génère la section des associés dans le document Word."""
    doc.add_heading('Liste des Associés', level=1)

    # Tableau des associés
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    # En-têtes
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Nom et Prénom'
    header_cells[1].text = 'Parts sociales'
    header_cells[2].text = 'Pourcentage'
    header_cells[3].text = 'Signature'

    # Données des associés
    for associe in associes:
        row_cells = table.add_row().cells
        row_cells[0].text = f"{associe['nom']} {associe['prenom']}"
        row_cells[1].text = str(associe['parts'])
        row_cells[2].text = f"{associe['pourcentage']}%"
        row_cells[3].text = ''  # Espace pour la signature


def convert_to_pdf(docx_path, pdf_path):
    """
    Convertit un fichier DOCX en PDF
    """
    word = None
    doc = None
    max_retries = 3
    retry_delay = 1

    for attempt in range(max_retries):
        try:
            import comtypes.client
            import pythoncom
            import time
            from pathlib import Path

            # Nettoyer les chemins
            docx_path = os.path.abspath(docx_path)
            pdf_path = os.path.abspath(pdf_path)

            # Vérifier que le fichier source existe
            if not os.path.exists(docx_path):
                raise FileNotFoundError(f"Le fichier source n'existe pas : {docx_path}")

            # S'assurer que le dossier de destination existe
            os.makedirs(os.path.dirname(pdf_path), exist_ok=True)

            # Initialiser COM pour le thread actuel
            pythoncom.CoInitialize()

            # Attendre que le fichier soit complètement écrit
            time.sleep(0.5)

            # Initialiser Word avec les paramètres optimaux
            word = comtypes.client.CreateObject('Word.Application')
            word.Visible = False
            word.DisplayAlerts = False
            word.Options.SaveNormalPrompt = False

            try:
                # Ouvrir le document
                doc = word.Documents.Open(docx_path)
                doc.Activate()

                # Sauvegarder en PDF avec les options optimales
                doc.ExportAsFixedFormat(
                    OutputFileName=pdf_path,
                    ExportFormat=17,  # wdExportFormatPDF = 17
                    OptimizeFor=0,    # wdExportOptimizeForPrint = 0
                    CreateBookmarks=0, # wdExportCreateNoBookmarks = 0
                    DocStructureTags=True,
                    BitmapMissingFonts=True,
                    UseISO19005_1=False
                )

                # Vérifier que le PDF a été créé
                if not os.path.exists(pdf_path):
                    raise Exception("Le PDF n'a pas été créé")

                return True

            except Exception as word_error:
                print(f"Erreur lors de l'opération Word (tentative {attempt + 1}/{max_retries}): {word_error}")
                if attempt == max_retries - 1:
                    raise

            finally:
                # Fermer proprement le document
                if doc:
                    try:
                        doc.Close(SaveChanges=False)
                        doc = None
                    except:
                        pass

        except Exception as e:
            print(f"Erreur lors de la conversion en PDF (tentative {attempt + 1}/{max_retries}): {e}")
            if attempt == max_retries - 1:
                # En dernier recours, essayer docx2pdf
                try:
                    from docx2pdf import convert
                    convert(docx_path, pdf_path)
                    return True
                except Exception as e2:
                    print(f"Erreur lors de la conversion alternative en PDF : {e2}")
                    raise

        finally:
            # Nettoyage des ressources Word
            if word:
                try:
                    word.Quit()
                    time.sleep(0.5)
                except:
                    pass

            # Libérer COM
            try:
                pythoncom.CoUninitialize()
            except:
                pass

            # Si ce n'est pas la dernière tentative, attendre avant de réessayer
            if attempt < max_retries - 1:
                time.sleep(retry_delay)

    return False

def save_to_database(data, docx_filename, pdf_filename):
    """
    Sauvegarde les informations du document dans la base de données
    """
    try:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()

        # Création de la table si elle n'existe pas
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS documents (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                type_societe TEXT NOT NULL,
                objet_ago TEXT NOT NULL,
                date_ago DATE NOT NULL,
                lieu_ago TEXT NOT NULL,
                heure_ago TIME NOT NULL,
                docx_filename TEXT NOT NULL,
                pdf_filename TEXT NOT NULL,
                date_creation TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')

        # Insertion des données
        cursor.execute('''
            INSERT INTO documents (
                type_societe, objet_ago, date_ago, lieu_ago,
                heure_ago, docx_filename, pdf_filename
            ) VALUES (?, ?, ?, ?, ?, ?, ?)
        ''', (
            data.get('type_societe', 'SARL'),
            data.get('objet_ago', ''),
            data.get('date_ago', ''),
            data.get('lieu_ago', ''),
            data.get('heure_ago', ''),
            docx_filename,
            pdf_filename
        ))

        conn.commit()

    except Exception as e:
        print(f"Erreur lors de la sauvegarde en base de données : {e}")
        raise
    finally:
        conn.close()

@app.route('/list_templates')
def list_templates():
    """Liste tous les modèles disponibles dans le répertoire templates."""
    try:
        templates_list = []
        # Parcourir le répertoire des templates
        for file in os.listdir(TEMPLATES_DIR):
            if file.endswith('.docx'):
                template_path = os.path.join(TEMPLATES_DIR, file)
                template_name = os.path.splitext(file)[0]
                templates_list.append({
                    'name': template_name,
                    'path': template_path,
                    'type': 'word'
                })

        return jsonify({
            'success': True,
            'templates': templates_list
        })
    except Exception as e:
        app.logger.error(f"Erreur lors de la lecture des templates : {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

def generate_filename(data, extension='pdf'):
    """
    Génère un nom de fichier au format AAAA_MM_JJ_Pv_Ago_Nom_de_Société
    """
    try:
        # Récupérer la date de l'AGO
        date_ago = datetime.strptime(data.get('date_ago', ''), '%Y-%m-%d')
        date_str = date_ago.strftime('%Y_%m_%d')
    except:
        # En cas d'erreur, utiliser la date actuelle
        date_str = datetime.now().strftime('%Y_%m_%d')

    # Récupérer et nettoyer le nom de la société
    nom_societe = data.get('nom_entreprise', '').strip()
    # Remplacer les caractères non autorisés par des underscores
    nom_societe = ''.join(c if c.isalnum() or c.isspace() else '_' for c in nom_societe)
    # Remplacer les espaces par des underscores et supprimer les underscores multiples
    nom_societe = '_'.join(filter(None, nom_societe.split('_')))

    return f"{date_str}_Pv_Ago_{nom_societe}.{extension}"

def get_default_values(type_societe='SARL'):
    """Génère des valeurs par défaut pour tester l'application."""
    current_date = datetime.now()
    default_values = {
        'type_societe': type_societe,
        'nom_entreprise': 'Société Test',
        'capital': '100000',
        'adresse': '123 Rue du Test, 75000 Paris',
        'date_assemblee': current_date.strftime('%Y-%m-%d'),
        'lieu_assemblee': 'Siège social',
        'heure_ago': '14:30',
        'objet_ago': 'Assemblée Générale Ordinaire Annuelle',
        'date_ago': current_date.strftime('%Y-%m-%d'),
        'lieu_ago': 'Siège social',
    }

    # Générer des associés par défaut selon le type de société
    if type_societe == 'SARL AU':
        default_values['associes'] = [{
            'nom': 'Doe',
            'prenom': 'John',
            'adresse': '123 Rue du Test, 75000 Paris',
            'cni': 'AA123456',
            'cni_validite': (current_date + timedelta(days=365)).strftime('%Y-%m-%d'),
            'cni_lieu': 'Paris',
            'email': 'john.doe@test.com',
            'telephone': '0123456789',
            'parts': 100,
            'pourcentage': 100.0
        }]
    else:  # SARL
        default_values['associes'] = [
            {
                'nom': 'Doe',
                'prenom': 'John',
                'adresse': '123 Rue du Test, 75000 Paris',
                'cni': 'AA123456',
                'cni_validite': (current_date + timedelta(days=365)).strftime('%Y-%m-%d'),
                'cni_lieu': 'Paris',
                'email': 'john.doe@test.com',
                'telephone': '0123456789',
                'parts': 60,
                'pourcentage': 60.0
            },
            {
                'nom': 'Smith',
                'prenom': 'Jane',
                'adresse': '456 Avenue du Test, 75000 Paris',
                'cni': 'BB789012',
                'cni_validite': (current_date + timedelta(days=365)).strftime('%Y-%m-%d'),
                'cni_lieu': 'Paris',
                'email': 'jane.smith@test.com',
                'telephone': '0987654321',
                'parts': 40,
                'pourcentage': 40.0
            }
        ]

    return default_values

def get_default_values():
    """Retourne les valeurs par défaut pour tous les champs du formulaire."""
    current_date = datetime.now()
    next_month = current_date + timedelta(days=30)
    previous_year = current_date.year - 1

    # Valeurs par défaut pour l'entreprise
    defaults = {
        # Section Entreprise
        'type_societe': 'SARL',  # ou 'SARL AU'
        'nom_entreprise': 'SOCIÉTÉ EXEMPLE',
        'capital': 100000,
        'adresse': '12, Rue Example, Quartier des Affaires, 20000 Casablanca',
        'if_number': '12345678',  # Format: 8 chiffres
        'rc_number': '123456',    # Format: 6 chiffres
        'ice_number': '123456789', # Format: 9 chiffres
        'cnss_number': '1234567',  # Format: 7 chiffres
        'tp_number': '12345678',   # Format: 8 chiffres

        # Section Associés (pour SARL par défaut)
        'associes': [
            {
                'nom': 'DUPONT',
                'prenom': 'Jean',
                'adresse': '15, Avenue de la Paix, 20100 Casablanca',
                'cni': 'BE123456',
                'cni_validite': (current_date + timedelta(days=365*5)).strftime('%Y-%m-%d'),  # Validité 5 ans
                'cni_lieu': 'Casablanca',
                'email': 'jean.dupont@example.com',
                'telephone': '0600000001',
                'parts': 600,
                'pourcentage': 60.0
            },
            {
                'nom': 'MARTIN',
                'prenom': 'Marie',
                'adresse': '25, Rue du Commerce, 20200 Casablanca',
                'cni': 'BE789012',
                'cni_validite': (current_date + timedelta(days=365*5)).strftime('%Y-%m-%d'),
                'cni_lieu': 'Casablanca',
                'email': 'marie.martin@example.com',
                'telephone': '0600000002',
                'parts': 400,
                'pourcentage': 40.0
            }
        ],

        # Section Détails AGO
        'exercice': previous_year,
        'date_ago': next_month.strftime('%Y-%m-%d'),
        'heure_ago': '10:00',
        'lieu_ago': 'au siège social de la société',
        'ordre_du_jour': [
            "Lecture du rapport de gestion du Gérant",
            f"Examen et approbation des comptes de l'exercice clos le 31 décembre {previous_year}",
            "Affectation du résultat de l'exercice",
            "Quitus au Gérant",
            "Questions diverses"
        ],

        # Informations financières
        'chiffre_affaires': 5000000,
        'resultat_net': 800000,
        'dividendes': 400000,
        'reserve_legale': 40000,  # 5% du résultat net
        'report_nouveau': 360000,  # Reste du résultat

        # Autres informations
        'date_creation': datetime(2020, 1, 1).strftime('%Y-%m-%d'),
        'forme_juridique': 'Société à Responsabilité Limitée',
        'duree': 99,
        'gerant': {
            'nom': 'DUPONT',
            'prenom': 'Jean',
            'fonction': 'Gérant',
            'date_nomination': datetime(2020, 1, 1).strftime('%Y-%m-%d')
        }
    }

    return defaults

if __name__ == '__main__':
    # Création des dossiers nécessaires au démarrage
    os.makedirs(os.path.join('static', 'templates'), exist_ok=True)
    os.makedirs(TEMP_DIR, exist_ok=True)
    app.run(debug=True)
